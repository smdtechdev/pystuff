import docx

doc = docx.Document('Automation\\files\\demo.docx')

print(doc)

print(doc.paragraphs[0].text)
print(doc.paragraphs[0].style)

p2 = doc.paragraphs[1]

print(p2.style)

# runs occur whenever a style changes in paragraph
print(p2.runs[0].text) # normal text
print(p2.runs[1].text) # bold text
print('is this run bold: ' + str(p2.runs[1].bold))
print(p2.runs[2].text) # normal text
print(p2.runs[3].text) # italic text
print('is this run italic: ' + str(p2.runs[3].italic))
#paragraph object p2 has 4 run objects

p2.runs[3].underline = True
p2.runs[3].text = 'Italic and Underline'

doc.save('Automation\\files\\demo2.docx') # text of the 4th run will be updated and underlined

doc2 = docx.Document() # create blank document object

doc2.add_paragraph('first para')
doc2.add_paragraph('second para')

p1 = doc2.paragraphs[0]
p1.add_run(' This is a new run and its bold')
p1.runs[1].bold = True

doc2.save('Automation\\files\\demo3.docx')

def getWordText(wordfile):
    doc = docx.Document(wordfile)
    allText = []
    for para in doc.paragraphs:
        allText.append(para.text)
    return '\n'.join(allText)

print(getWordText('Automation\\files\\demo.docx'))
        